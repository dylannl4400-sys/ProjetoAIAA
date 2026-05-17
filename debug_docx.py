# -*- coding: utf-8 -*-
import sys, io, re
from pathlib import Path
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

try:
    from docx import Document
except ImportError:
    print("ERRO: pip install python-docx")
    sys.exit(1)

CAMPOS = [
    "nome_parte_a", "nome_parte_b", "tribunal", "juiz", "juizo",
    "referencia_processo", "descricao_factos", "pedido", "valor",
    "junta", "nome_escritorio", "base_legal", "template",
    "fundamentos",
]

def analisar(caminho):
    doc = Document(caminho)
    print(f"\n{'='*60}")
    print(f"FICHEIRO: {caminho}")
    print(f"{'='*60}\n")

    marcadores = set()

    print("--- PARAGRAFOS ---")
    for i, para in enumerate(doc.paragraphs):
        texto = para.text
        if not texto.strip():
            continue
        found = re.findall(r'\{\{[^}]+\}\}', texto)
        marcadores.update(found)
        if found:
            print(f"\n[P{i}] TEXTO: {repr(texto)}")
            print(f"      MARCADORES: {found}")
            print(f"      RUNS ({len(para.runs)}):")
            for j, run in enumerate(para.runs):
                print(f"        run[{j}]: {repr(run.text)}")

    print("\n--- TABELAS ---")
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, para in enumerate(cell.paragraphs):
                    texto = para.text
                    if not texto.strip():
                        continue
                    found = re.findall(r'\{\{[^}]+\}\}', texto)
                    marcadores.update(found)
                    if found:
                        print(f"\n[T{ti}R{ri}C{ci}P{pi}] {repr(texto)}")
                        for j, run in enumerate(para.runs):
                            print(f"  run[{j}]: {repr(run.text)}")

    print(f"\n{'='*60}")
    print(f"MARCADORES NO DOCUMENTO ({len(marcadores)}):")
    for m in sorted(marcadores):
        chave = m[2:-2]
        match = "OK" if chave in CAMPOS else "SEM MATCH no formulario"
        print(f"  {m}  --> {match}")

    if not marcadores:
        print("  NENHUM MARCADOR {{...}} ENCONTRADO!")
        print("  Verifica se o Word usa { e } normais (nao curvos).")

    print(f"\n{'='*60}")
    print("SUBSTITUICOES QUE O CODIGO VAI TENTAR:")
    for campo in CAMPOS:
        print(f"  {{{{{campo}}}}}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        pasta = Path("pipeline/legal_docs")
        docs = list(pasta.glob("template_*.docx"))
        if not docs:
            print("Uso: python debug_docx.py <ficheiro.docx>")
            sys.exit(1)
        print(f"A analisar: {docs[-1]}")
        analisar(str(docs[-1]))
    else:
        analisar(sys.argv[1])
