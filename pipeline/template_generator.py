# """
# template_generator.py

# Geração de peças processuais a partir dos templates .docx reais.

# Marcadores confirmados por template:
#     Contestação:  XXXXX (tribunal), XXXXXX (juízo), XXXX (juiz/proc),
#                   (Nome da parte), XXXX. (fundamentos), XXXXX (valor)
#     Petição:      XXXXXXXXXX (autor), XXXXXXXXXXX (réu/junta),
#                   XXXXXXXXXXX (fundamentos), XXXXXX (valor)
#     Requerimento: XXXXX (tribunal), (Nome da parte), XXXXXXX (réu),
#                   XXXXXXXXX (base legal), XXXXXXXX (fundamentos),
#                   XXXX. (pedido), XXXXX (valor), XXXXXX (junta)
# """

# import os
# import re
# import tempfile
# from pathlib import Path

# try:
#     from docx import Document
#     HAS_DOCX = True
# except ImportError:
#     HAS_DOCX = False

# TEMPLATES_DIR = Path(__file__).parent / "templates"

# TEMPLATE_MAP = {
#     "contestacao":  "contestacao.docx",
#     "peticao":      "peticao.docx",
#     "requerimento": "requerimento.docx",
# }


# class TemplateGenerator:

#     def __init__(self, llm=None):
#         self.llm = llm
#         if not HAS_DOCX:
#             raise ImportError("Instala python-docx: pip install python-docx")

#     def gerar(self, tipo: str, dados: dict, output_path: str = None) -> str:
#         template_file = TEMPLATE_MAP.get(tipo)
#         if not template_file:
#             raise ValueError(f"Tipo '{tipo}' não reconhecido.")

#         template_path = TEMPLATES_DIR / template_file
#         if not template_path.exists():
#             raise FileNotFoundError(f"Template não encontrado: {template_path}")

#         conteudo = self._gerar_conteudo_llm(tipo, dados)
#         doc      = Document(str(template_path))
#         subs     = self._construir_substituicoes(tipo, dados, conteudo)
#         self._preencher(doc, subs)

#         if not output_path:
#             tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
#             tmp.close()
#             output_path = tmp.name

#         doc.save(output_path)
#         return output_path

#     # ------------------------------------------------------------------
#     # Substituições por template
#     # ------------------------------------------------------------------

#     def _construir_substituicoes(self, tipo: str, dados: dict, conteudo: dict) -> list:
#         """
#         Devolve lista de (marcador_exacto, valor) por ordem de prioridade.
#         Mais específicos primeiro para evitar substituições parciais.
#         """
#         tribunal    = dados.get("tribunal",            "")
#         parte_a     = dados.get("nome_parte_a",        "")
#         parte_b     = dados.get("nome_parte_b",        "")
#         processo    = dados.get("referencia_processo", "")
#         valor       = dados.get("valor",               "")
#         junta       = dados.get("junta",               "")
#         advogado    = dados.get("nome_escritorio",     "")
#         data        = dados.get("local_data",          "")
#         fundamentos = conteudo.get("fundamentos",      dados.get("descricao_factos", ""))
#         pedido      = conteudo.get("pedido",           dados.get("pedido", ""))
#         base_legal  = dados.get("base_legal",          "")

#         if tipo == "contestacao":
#             return [
#                 # Processo — mais específico primeiro
#                 ("XXXX/XX.XXXXXXX",      processo),
#                 ("XXXXXX de XXXX",       dados.get("juizo", "")),
#                 # Tribunal
#                 ("XXXXX",                tribunal),
#                 # Partes
#                 ("(Nome da parte)",      parte_a),
#                 # Fundamentos — parágrafo [9] "XXXX."
#                 ("XXXX.",                fundamentos),
#                 # Pedido — parágrafo [14]
#                 ("Julgar-se como não provado,;", pedido),
#                 # Prova e documentos
#                 ("XX documentos",        junta or "ver lista junta"),
#                 # Valor
#                 ("XXXXX",               valor or "a fixar"),
#             ]

#         elif tipo == "peticao":
#             return [
#                 # Autor — parágrafo [8] "XXXXXXXXXX, vêm..."
#                 # Substituímos só o marcador, mantendo o resto do texto
#                 ("XXXXXXXXXX",           parte_a),
#                 # Réu — parágrafo [11]
#                 ("XXXXXXXXXXX",          parte_b),
#                 # Fundamentos — parágrafo [14]
#                 # (mesmo marcador que réu mas parágrafo diferente — ordem importa)
#                 # Valor
#                 ("XXXXXX",              valor or "a fixar"),
#                 # Junta
#                 ("XXXXXXXXXXX",          junta or "ver lista junta"),
#                 # Processo/referência no título
#                 ("XXXXXXX",             processo),
#             ]

#         elif tipo == "requerimento":
#             return [
#                 # Tribunal
#                 ("XXXXX",               tribunal),
#                 # Partes
#                 ("(Nome da parte)",     parte_a),
#                 ("XXXXXXX",             parte_b),
#                 # Base legal — parágrafo [6] "Nos termos do XXXXXXXXX"
#                 ("XXXXXXXXX",           base_legal or "artigo aplicável"),
#                 # Fundamentos — parágrafo [7] "XXXXXXXX"
#                 ("XXXXXXXX",            fundamentos),
#                 # Pedido — parágrafo [9] "XXXX."
#                 ("XXXX.",               pedido),
#                 # Valor e junta
#                 ("XXXXX",               valor or "a fixar"),
#                 ("XXXXXX",              junta or "ver lista junta"),
#             ]

#         return []

#     # ------------------------------------------------------------------
#     # Preenchimento do documento
#     # ------------------------------------------------------------------

#     def _preencher(self, doc, substituicoes: list):
#         """Aplica substituições a todos os parágrafos e tabelas."""
#         for para in doc.paragraphs:
#             self._sub_paragrafo(para, substituicoes)
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     for para in cell.paragraphs:
#                         self._sub_paragrafo(para, substituicoes)

#     def _sub_paragrafo(self, para, substituicoes: list):
#         """
#         Substitui marcadores num parágrafo.
#         Reconstrói o texto completo para lidar com runs fragmentados.
#         """
#         # Texto completo do parágrafo
#         texto = para.text
#         if not texto.strip():
#             return

#         # Verificar se há algum marcador neste parágrafo
#         tem_marcador = any(old in texto for old, _ in substituicoes)
#         if not tem_marcador:
#             return

#         # Aplicar substituições ao texto completo
#         texto_novo = texto
#         for old, new in substituicoes:
#             if old in texto_novo:
#                 texto_novo = texto_novo.replace(old, str(new) if new else "")

#         if texto_novo == texto:
#             return

#         # Reescrever o parágrafo — colocar o texto no primeiro run
#         # e limpar os restantes
#         if para.runs:
#             # Guardar formatação do primeiro run
#             primeiro = para.runs[0]
#             primeiro.text = texto_novo
#             # Limpar runs seguintes
#             for run in para.runs[1:]:
#                 run.text = ""
#         else:
#             # Sem runs — adicionar um
#             para.add_run(texto_novo)

#     # ------------------------------------------------------------------
#     # LLM
#     # ------------------------------------------------------------------

#     def _gerar_conteudo_llm(self, tipo: str, dados: dict) -> dict:
#         if not self.llm:
#             return {
#                 "fundamentos": dados.get("descricao_factos", ""),
#                 "pedido":      dados.get("pedido", ""),
#             }

#         nomes = {
#             "contestacao":  "Contestação",
#             "peticao":      "Petição Inicial de Intimação para Protecção de Direitos",
#             "requerimento": "Requerimento",
#         }

#         prompt = f"""És um advogado português especializado em direito processual civil e administrativo.
# Gera os fundamentos jurídicos para uma {nomes.get(tipo, tipo)} com base nos seguintes dados:

# Requerente/Autor: {dados.get('nome_parte_a', '')}
# Requerido/Réu: {dados.get('nome_parte_b', '')}
# Tribunal: {dados.get('tribunal', '')}
# Factos: {dados.get('descricao_factos', '')}
# Pedido do cliente: {dados.get('pedido', '')}

# Responde APENAS com este formato (sem introdução nem conclusão):

# FUNDAMENTOS:
# [fundamentos jurídicos detalhados com referência a artigos de lei]

# PEDIDO:
# [formulação jurídica precisa do pedido]

# Usa linguagem jurídica formal portuguesa. Sê concreto e preciso."""

#         try:
#             resultado = self.llm.generate(prompt)
#             return {
#                 "fundamentos": self._extrair(resultado, "FUNDAMENTOS"),
#                 "pedido":      self._extrair(resultado, "PEDIDO"),
#             }
#         except Exception:
#             return {
#                 "fundamentos": dados.get("descricao_factos", ""),
#                 "pedido":      dados.get("pedido", ""),
#             }

#     def _extrair(self, texto: str, secao: str) -> str:
#         m = re.search(rf"{secao}:\s*\n(.*?)(?=\n[A-ZÇÃÕ]{{3,}}:|$)", texto, re.DOTALL)
#         return m.group(1).strip() if m else texto.strip()


# # ------------------------------------------------------------------
# # Função de conveniência
# # ------------------------------------------------------------------

# def gerar_peca_processual(tipo: str, dados: dict, llm=None) -> str:
#     if tipo == "cessacao":
#         from document_generator import DocumentGenerator, CasoDespedimento
#         gen  = DocumentGenerator(None, llm)
#         caso = CasoDespedimento(
#             nome_trabalhador         = dados.get("nome_parte_a", ""),
#             nome_empregador          = dados.get("nome_parte_b", ""),
#             descricao_factos         = dados.get("descricao_factos", ""),
#             data_efeitos             = dados.get("local_data", ""),
#             motivo_cessacao          = dados.get("motivo_cessacao", "Justa Causa Disciplinar"),
#             referencia_processo      = dados.get("referencia_processo", ""),
#             nif_trabalhador          = dados.get("nif_parte_a", ""),
#             nif_empregador           = dados.get("nif_parte_b", ""),
#             morada_trabalhador       = dados.get("morada_parte_a", ""),
#             morada_empregador        = dados.get("morada_parte_b", ""),
#             nome_escritorio          = dados.get("nome_escritorio", ""),
#             representante_empregador = dados.get("nome_parte_b", ""),
#             cargo_representante      = "",
#             categoria_profissional   = "",
#             data_admissao            = "",
#             numero_contrato          = "",
#             artigo_cessacao          = "351.º e 357.º",
#             local_data               = dados.get("local_data", ""),
#         )
#         return gen.gerar(caso)

#     gen = TemplateGenerator(llm=llm)
#     return gen.gerar(tipo, dados)

# """
# template_generator.py

# Dois modos de geração:
#     - Modo simples: LLM usa apenas os dados do formulário
#     - Modo RAG: LLM usa dados + chunks de acórdãos indexados no ChromaDB
# """

# import os, re, tempfile
# from pathlib import Path

# try:
#     from docx import Document
#     HAS_DOCX = True
# except ImportError:
#     HAS_DOCX = False

# TEMPLATES_DIR = Path(__file__).parent / "legal_docs"

# TEMPLATE_MAP = {
#     "contestacao":  "contestacao.docx",
#     "peticao":      "peticao.docx",
#     "requerimento": "requerimento.docx",
# }


# class TemplateGenerator:

#     def __init__(self, llm=None, store=None):
#         self.llm   = llm
#         self.store = store  # se fornecido activa modo RAG
#         if not HAS_DOCX:
#             raise ImportError("pip install python-docx")

#     def gerar(self, tipo: str, dados: dict, output_path: str = None) -> str:
#         template_file = TEMPLATE_MAP.get(tipo)
#         if not template_file:
#             raise ValueError(f"Tipo '{tipo}' não reconhecido. Disponíveis: {list(TEMPLATE_MAP)}")

#         template_path = TEMPLATES_DIR / template_file
#         if not template_path.exists():
#             raise FileNotFoundError(f"Template não encontrado: {template_path}")

#         conteudo = self._gerar_conteudo(tipo, dados)
#         doc      = Document(str(template_path))
#         subs     = self._substituicoes(tipo, dados, conteudo)
#         self._preencher(doc, subs)

#         if not output_path:
#             tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
#             tmp.close()
#             output_path = tmp.name

#         doc.save(output_path)
#         return output_path

#     # ------------------------------------------------------------------
#     # Conteúdo — simples ou RAG
#     # ------------------------------------------------------------------

#     def _gerar_conteudo(self, tipo: str, dados: dict) -> dict:
#         """Gera fundamentos e pedido — com ou sem RAG consoante store disponível."""
#         if not self.llm:
#             return {
#                 "fundamentos": dados.get("descricao_factos", ""),
#                 "pedido":      dados.get("pedido", ""),
#                 "fonte":       "formulário",
#             }

#         # Buscar chunks relevantes se store disponível
#         contexto    = ""
#         acordaos    = []
#         if self.store:
#             query  = f"{dados.get('descricao_factos','')} {dados.get('pedido','')}"
#             chunks = self.store.search(query, n=6)
#             chunks = [c for c in chunks if c["score"] > 0.50]
#             if chunks:
#                 partes = []
#                 vistos = set()
#                 for c in chunks:
#                     proc = c["metadata"].get("processo", "")
#                     if proc and proc not in vistos:
#                         vistos.add(proc)
#                         acordaos.append(proc)
#                     partes.append(
#                         f"[Acórdão {proc or 'desconhecido'} — {c['metadata'].get('section','')}]\n{c['text']}"
#                     )
#                 contexto = "\n\n---\n\n".join(partes)

#         prompt = self._prompt(tipo, dados, contexto)

#         try:
#             resultado = self.llm.generate(prompt)
#             return {
#                 "fundamentos": self._extrair(resultado, "FUNDAMENTOS"),
#                 "pedido":      self._extrair(resultado, "PEDIDO"),
#                 "fonte":       "rag" if contexto else "llm",
#                 "acordaos":    acordaos,
#             }
#         except Exception as e:
#             return {
#                 "fundamentos": dados.get("descricao_factos", ""),
#                 "pedido":      dados.get("pedido", ""),
#                 "fonte":       "fallback",
#                 "erro":        str(e),
#             }

#     def _prompt(self, tipo: str, dados: dict, contexto: str) -> str:
#         nomes = {
#             "contestacao":  "Contestação",
#             "peticao":      "Petição Inicial de Intimação para Protecção de Direitos",
#             "requerimento": "Requerimento",
#         }
#         secao_contexto = f"""
# JURISPRUDÊNCIA RELEVANTE (acórdãos indexados):
# {contexto}

# Com base nesta jurisprudência, fundamenta juridicamente a peça processual.
# Cita os acórdãos relevantes quando adequado (ex: "conforme Acórdão do TRE, proc. XXXX").
# """ if contexto else """
# Não há jurisprudência indexada. Fundamenta com base na lei geral portuguesa aplicável.
# """

#         return f"""És um advogado português especializado em direito processual civil e administrativo.
# Gera os fundamentos jurídicos para uma {nomes.get(tipo, tipo)}.

# DADOS DO CASO:
# Requerente/Autor: {dados.get('nome_parte_a', '')}
# Requerido/Réu:    {dados.get('nome_parte_b', '')}
# Tribunal:         {dados.get('tribunal', '')}
# Factos:           {dados.get('descricao_factos', '')}
# Pedido do cliente: {dados.get('pedido', '')}
# {secao_contexto}
# Responde APENAS com este formato (sem introdução nem conclusão):

# FUNDAMENTOS:
# [fundamentos jurídicos detalhados, linguagem formal portuguesa]

# PEDIDO:
# [formulação jurídica precisa do pedido ao tribunal]"""

#     def _extrair(self, texto: str, secao: str) -> str:
#         m = re.search(rf"{secao}:\s*\n(.*?)(?=\n[A-ZÇÃÕÁÉÍÓÚ]{{3,}}:|$)", texto, re.DOTALL)
#         return m.group(1).strip() if m else texto.strip()

#     # ------------------------------------------------------------------
#     # Substituições por template
#     # ------------------------------------------------------------------

#     def _substituicoes(self, tipo: str, dados: dict, conteudo: dict) -> list:
#         tribunal    = dados.get("tribunal",            "")
#         parte_a     = dados.get("nome_parte_a",        "")
#         parte_b     = dados.get("nome_parte_b",        "")
#         processo    = dados.get("referencia_processo", "")
#         valor       = dados.get("valor",               "a fixar")
#         junta       = dados.get("junta",               "ver lista junta")
#         base_legal  = dados.get("base_legal",          "artigo aplicável")
#         fundamentos = conteudo.get("fundamentos", dados.get("descricao_factos", ""))
#         pedido      = conteudo.get("pedido",      dados.get("pedido", ""))

#         if tipo == "contestacao":
#             return [
#                 ("XXXX/XX.XXXXXXX",             processo),
#                 ("Juízo XXXXXX de XXXX",               dados.get("juizo", "")),
#                 ("Juiz X",                             dados.get("juiz", "")),
#                 ("XXXXX",                        tribunal),
#                 ("(Nome da parte)",              parte_a),
#                 ("XXXX.",                        fundamentos),
#                 ("Julgar-se como não provado,;", pedido),
#                 ("XX documentos",                junta),
#                 ("XXXXX",                        valor),
#             ]
#         elif tipo == "peticao":
#             return [
#                 ("{{PARTE_A}}",  parte_a),
#                 ("{{PARTE_B}}", parte_b),
#                 ("{{VALOR}}",      valor),
#                 ("XXXX",                        fundamentos),
#                 ("{{JUNTA}}", junta),
#                 ("{{REF}}",     processo),
#             ]
#         elif tipo == "requerimento":
#             return [
#                 ("Tribunal Judicial da Comarca de XXXXX",       tribunal),
#                 ("(Nome da parte), xxxxxxx,", parte_a),
#                 ("{{PARTE_B}}",     parte_b),
#                 ("{{BASE_LEGAL}}",   base_legal),
#                 ("XXXXXXXX",    fundamentos),
#                 ("{{PEDIDO}}",       pedido),
#                 ("{{VALOR}}",       valor),
#                 ("{{JUNTA}}",      junta),
#             ]
#         return []

#     # ------------------------------------------------------------------
#     # Preenchimento
#     # ------------------------------------------------------------------

#     def _preencher(self, doc, substituicoes: list):
#         for para in doc.paragraphs:
#             self._sub_para(para, substituicoes)
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     for para in cell.paragraphs:
#                         self._sub_para(para, substituicoes)

#     def _sub_para(self, para, substituicoes: list):
#         texto = para.text
#         if not texto.strip():
#             return
#         if not any(old in texto for old, _ in substituicoes):
#             return
#         texto_novo = texto
#         for old, new in substituicoes:
#             if old in texto_novo:
#                 texto_novo = texto_novo.replace(old, str(new) if new else "")
#         if texto_novo == texto:
#             return
#         if para.runs:
#             para.runs[0].text = texto_novo
#             for run in para.runs[1:]:
#                 run.text = ""
#         else:
#             para.add_run(texto_novo)


# # ------------------------------------------------------------------
# # Função de conveniência
# # ------------------------------------------------------------------

# def gerar_peca_processual(tipo: str, dados: dict, llm=None, store=None) -> str:
#     """
#     Gera uma peça processual.

#     Se store=None  → modo simples (só dados do formulário + LLM)
#     Se store!=None → modo RAG (dados + acórdãos indexados + LLM)
#     """
#     gen = TemplateGenerator(llm=llm, store=store)
#     resultado = gen._gerar_conteudo(tipo, dados)

#     # Log para debug
#     fonte = resultado.get("fonte", "?")
#     acordaos = resultado.get("acordaos", [])
#     if acordaos:
#         print(f"[TemplateGenerator] RAG: {len(acordaos)} acórdãos usados: {acordaos}")
#     else:
#         print(f"[TemplateGenerator] Modo: {fonte}")

#     return gen.gerar(tipo, dados)


"""
template_generator.py

Dois modos de geração:
    - Modo simples: LLM usa apenas os dados do formulário
    - Modo RAG: LLM usa dados + chunks de acórdãos indexados no ChromaDB
"""

import os, re, tempfile
from pathlib import Path

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

TEMPLATES_DIR = Path(__file__).parent / "legal_docs"

TEMPLATE_MAP = {
    "contestacao":  "contestacao.docx",
    "peticao":      "peticao.docx",
    "requerimento": "requerimento.docx",
    "cessacao":     "carta_cessacao_v1.docx",
}


class TemplateGenerator:

    def __init__(self, llm=None, store=None):
        self.llm   = llm
        self.store = store  # se fornecido activa modo RAG
        if not HAS_DOCX:
            raise ImportError("pip install python-docx")

    def gerar(self, tipo_ou_ficheiro: str, dados: dict, output_path: str = None, conteudo: dict = None) -> str:
        # Se contiver .docx, assume que é o nome do ficheiro direto
        if tipo_ou_ficheiro.endswith(".docx"):
            template_file = tipo_ou_ficheiro
            tipo = tipo_ou_ficheiro.split(".")[0] # fallback para o tipo
        else:
            template_file = TEMPLATE_MAP.get(tipo_ou_ficheiro)
            tipo = tipo_ou_ficheiro

        if not template_file:
            raise ValueError(f"Template '{tipo_ou_ficheiro}' não reconhecido.")

        template_path = TEMPLATES_DIR / template_file
        if not template_path.exists():
            raise FileNotFoundError(f"Template não encontrado: {template_path}")

        if not conteudo:
            conteudo = self._gerar_conteudo(tipo, dados)
        doc      = Document(str(template_path))
        subs     = self._substituicoes(tipo, dados, conteudo)
        self._preencher(doc, subs)

        if not output_path:
            tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            tmp.close()
            output_path = tmp.name

        doc.save(output_path)
        return output_path

    # ------------------------------------------------------------------
    # Conteúdo — simples ou RAG
    # ------------------------------------------------------------------

    def _gerar_conteudo(self, tipo: str, dados: dict) -> dict:
        """Gera fundamentos e pedido — com ou sem RAG consoante store disponível."""
        if not self.llm:
            return {
                "fundamentos": dados.get("descricao_factos", ""),
                "pedido":      dados.get("pedido", ""),
                "fonte":       "formulário",
            }

        # Buscar chunks relevantes se store disponível
        contexto    = ""
        acordaos    = []
        if self.store:
            query  = f"{dados.get('descricao_factos','')} {dados.get('pedido','')}"
            chunks = self.store.search(query, n=6)
            chunks = [c for c in chunks if c["score"] > 0.50]
            if chunks:
                partes = []
                vistos = set()
                for c in chunks:
                    proc = c["metadata"].get("processo", "")
                    if proc and proc not in vistos:
                        vistos.add(proc)
                        acordaos.append(proc)
                    partes.append(
                        f"[Acórdão {proc or 'desconhecido'} — {c['metadata'].get('section','')}]\n{c['text']}"
                    )
                contexto = "\n\n---\n\n".join(partes)

        prompt = self._prompt(tipo, dados, contexto)

        try:
            resultado = self.llm.generate(prompt)
            return {
                "fundamentos": self._extrair(resultado, "FUNDAMENTOS"),
                "pedido":      self._extrair(resultado, "PEDIDO"),
                "fonte":       "rag" if contexto else "llm",
                "acordaos":    acordaos,
            }
        except Exception as e:
            return {
                "fundamentos": dados.get("descricao_factos", ""),
                "pedido":      dados.get("pedido", ""),
                "fonte":       "fallback",
                "erro":        str(e),
            }

    def _prompt(self, tipo: str, dados: dict, contexto: str) -> str:
        nomes = {
            "contestacao":  "Contestação",
            "peticao":      "Petição Inicial de Intimação para Protecção de Direitos",
            "requerimento": "Requerimento",
        }
        secao_contexto = f"""
JURISPRUDÊNCIA RELEVANTE (acórdãos indexados):
{contexto}

Com base nesta jurisprudência, fundamenta juridicamente a peça processual.
Cita os acórdãos relevantes quando adequado (ex: "conforme Acórdão do TRE, proc. XXXX").
""" if contexto else """
Não há jurisprudência indexada. Fundamenta com base na lei geral portuguesa aplicável.
"""

        return f"""És um advogado português especializado em direito processual civil e administrativo.
Gera os fundamentos jurídicos para uma {nomes.get(tipo, tipo)}.

DADOS DO CASO:
Requerente/Autor: {dados.get('nome_parte_a', '')}
Requerido/Réu:    {dados.get('nome_parte_b', '')}
Tribunal:         {dados.get('tribunal', '')}
Factos:           {dados.get('descricao_factos', '')}
Pedido do cliente: {dados.get('pedido', '')}
{secao_contexto}
Responde APENAS com este formato (sem introdução nem conclusão):

FUNDAMENTOS:
[fundamentos jurídicos detalhados, linguagem formal portuguesa]

PEDIDO:
[formulação jurídica precisa do pedido ao tribunal]"""

    def _extrair(self, texto: str, secao: str) -> str:
        m = re.search(rf"{secao}:\s*\n(.*?)(?=\n[A-ZÇÃÕÁÉÍÓÚ]{{3,}}:|$)", texto, re.DOTALL)
        return m.group(1).strip() if m else texto.strip()

    # ------------------------------------------------------------------
    # Substituições por template
    # ------------------------------------------------------------------

    # def _substituicoes(self, tipo: str, dados: dict, conteudo: dict) -> list:
    #     tribunal    = dados.get("tribunal",            "")
    #     parte_a     = dados.get("nome_parte_a",        "")
    #     parte_b     = dados.get("nome_parte_b",        "")
    #     processo    = dados.get("referencia_processo", "")
    #     valor       = dados.get("valor",               "a fixar")
    #     junta       = dados.get("junta",               "ver lista junta")
    #     base_legal  = dados.get("base_legal",          "artigo aplicável")
    #     fundamentos = conteudo.get("fundamentos", dados.get("descricao_factos", ""))
    #     pedido      = conteudo.get("pedido",      dados.get("pedido", ""))

    #     if tipo == "contestacao":
    #         return [
    #             ("XXXX/XX.XXXXXXX",             processo),
    #             ("XXXXXX de XXXX",               dados.get("juizo", "")),
    #             ("XXXXX",                        tribunal),
    #             ("(Nome da parte)",              parte_a),
    #             ("XXXX.",                        fundamentos),
    #             ("Julgar-se como não provado,;", pedido),
    #             ("XX documentos",                junta),
    #             ("XXXXX",                        valor),
    #         ]
    #     elif tipo == "peticao":
    #         return [
    #             ("XXXXXXXXXX",  parte_a),
    #             ("XXXXXXXXXXX", parte_b),
    #             ("XXXXXX",      valor),
    #             ("XXXXXXXXXXX", junta),
    #             ("XXXXXXX",     processo),
    #         ]
    #     elif tipo == "requerimento":
    #         return [
    #             ("XXXXX",       tribunal),
    #             ("(Nome da parte)", parte_a),
    #             ("XXXXXXX",     parte_b),
    #             ("XXXXXXXXX",   base_legal),
    #             ("XXXXXXXX",    fundamentos),
    #             ("XXXX.",       pedido),
    #             ("XXXXX",       valor),
    #             ("XXXXXX",      junta),
    #         ]
    #     return []
    
    def _substituicoes(self, tipo: str, dados: dict, conteudo: dict) -> list:
        tribunal    = dados.get("tribunal",            "")
        parte_a     = dados.get("nome_parte_a",        "")
        parte_b     = dados.get("nome_parte_b",        "")
        processo    = dados.get("referencia_processo", "")
        valor       = dados.get("valor",               "a fixar")
        junta       = dados.get("junta",               "ver lista junta")
        base_legal  = dados.get("base_legal",          "artigo aplicável")
        fundamentos = conteudo.get("fundamentos", dados.get("descricao_factos", ""))
        pedido      = conteudo.get("pedido",      dados.get("pedido", ""))

        if tipo == "contestacao":
            res = [
                ("{{processo}}",             processo),
                ("{{Juízo juízo}}",               dados.get("juizo", "")),
                ("{{Juiz juiz}}",                             dados.get("juiz", "")),
                ("{{Tribunal Judicial da Comarca de XXXXX}}",                        tribunal),
                ("{{(Nome da parte)}}",              parte_a),
                ("{{fun}}",                        fundamentos),
                ("{{Julgar-se como não provado,;}}", pedido),
                ("{{XX documentos}}",                junta),
                ("{{valor}}",                        valor),
            ]
        elif tipo == "peticao":
            res = [
                ("{{PARTE_A}}",  parte_a),
                ("{{PARTE_B}}", parte_b),
                ("{{VALOR}}",      valor),
                ("{{XXXX}}",                        fundamentos),
                ("{{JUNTA}}", junta),
                ("{{REF}}",     processo),
            ]
        elif tipo == "requerimento":
            res = [
                ("{{Tribunal Judicial da Comarca de XXXXX}}",       tribunal),
                ("{{(Nome da parte), xxxxxxx}}", parte_a),
                ("{{PARTE_B}}",     parte_b),
                ("{{BASE_LEGAL}}",   base_legal),
                ("{{XXXXXXXX}}",    fundamentos),
                ("{{PEDIDO}}",       pedido),
                ("{{VALOR}}",       valor),
                ("{{JUNTA}}",      junta),
            ]
        elif tipo == "cessacao":
            res = [
                ("{{NOME_PARTE_A}}",          dados.get("nome_parte_a", "")),
                ("{{NIF_PARTE_A}}",           dados.get("nif_parte_a", "")),
                ("{{MORADA_PARTE_A}}",        dados.get("morada_parte_a", "")),
                ("{{NOME_PARTE_B}}",          dados.get("nome_parte_b", "")),
                ("{{NIF_PARTE_B}}",           dados.get("nif_parte_b", "")),
                ("{{MORADA_PARTE_B}}",        dados.get("morada_parte_b", "")),
                ("{{PROCESSO}}",              dados.get("processo", "")),
                ("{{ADVOGADO}}",              dados.get("advogado", "")),
                ("{{REPRESENTANTE}}",         dados.get("representante", "")),
                ("{{CATEGORIA}}",             dados.get("categoria", "")),
                ("{{DATA_ADMISSAO}}",         dados.get("data_admissao", "")),
                ("{{MOTIVO}}",                dados.get("motivo", "")),
                ("{{DATA_EFEITOS}}",          dados.get("data_efeitos", "")),
                ("{{LOCAL_DATA}}",            dados.get("local_data", "")),
                ("{{DESCRICAO_FACTOS}}",      dados.get("descricao_factos", "")),
            ]
        else:
            res = []

        # Adicionar substituições genéricas para qualquer campo do formulário {{campo}}
        for k, v in dados.items():
            res.append((f"{{{{{k}}}}}", str(v) if v else ""))
        
        # Adicionar fundamentos e pedido como {{fundamentos}} e {{pedido}}
        res.append(("{{fundamentos}}", fundamentos))
        res.append(("{{pedido}}", pedido))

        return res

    # ------------------------------------------------------------------
    # Preenchimento
    # ------------------------------------------------------------------

    def _preencher(self, doc, substituicoes: list):
        for para in doc.paragraphs:
            self._sub_para(para, substituicoes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._sub_para(para, substituicoes)

    def _sub_para(self, para, substituicoes: list):
        """
        Substitui marcadores preservando os estilos de cada run.

        Estratégia:
        1. Tentar substituir dentro de cada run individualmente
           (funciona quando o marcador está num único run)
        2. Se não resultar, reconstruir o texto completo no primeiro run
           preservando a formatação do run original
        """
        texto_completo = para.text
        if not texto_completo.strip():
            return
        if not any(old in texto_completo for old, _ in substituicoes):
            return

        # Tentativa 1: substituir dentro de cada run individualmente
        # Preserva formatação quando o marcador está num só run
        alterado = False
        for run in para.runs:
            t = run.text
            for old, new in substituicoes:
                if old in t:
                    t = t.replace(old, str(new) if new else "")
                    alterado = True
            run.text = t

        if alterado:
            return

        # Tentativa 2: o marcador está fragmentado entre vários runs
        # Reconstrói o parágrafo preservando o estilo do primeiro run
        texto_novo = texto_completo
        for old, new in substituicoes:
            if old in texto_novo:
                texto_novo = texto_novo.replace(old, str(new) if new else "")

        if texto_novo == texto_completo:
            return

        # Guardar formatação do primeiro run
        if para.runs:
            primeiro_run = para.runs[0]
            # Copiar propriedades de formatação
            bold      = primeiro_run.bold
            italic    = primeiro_run.italic
            underline = primeiro_run.underline
            font_name = primeiro_run.font.name
            font_size = primeiro_run.font.size
            color     = primeiro_run.font.color.rgb if primeiro_run.font.color and primeiro_run.font.color.type else None

            # Limpar todos os runs
            primeiro_run.text = texto_novo
            for run in para.runs[1:]:
                run.text = ""

            # Restaurar formatação
            primeiro_run.bold      = bold
            primeiro_run.italic    = italic
            primeiro_run.underline = underline
            if font_name:
                primeiro_run.font.name = font_name
            if font_size:
                primeiro_run.font.size = font_size
        else:
            para.add_run(texto_novo)


# ------------------------------------------------------------------
# Função de conveniência
# ------------------------------------------------------------------

def gerar_peca_processual(tipo: str, dados: dict, llm=None, store=None) -> str:
    """
    Gera uma peça processual.

    Se store=None  → modo simples (só dados do formulário + LLM)
    Se store!=None → modo RAG (dados + acórdãos indexados + LLM)
    """
    gen = TemplateGenerator(llm=llm, store=store)
    resultado = gen._gerar_conteudo(tipo, dados)

    # Log para debug
    fonte = resultado.get("fonte", "?")
    acordaos = resultado.get("acordaos", [])
    if acordaos:
        print(f"[TemplateGenerator] RAG: {len(acordaos)} acórdãos usados: {acordaos}")
    else:
        print(f"[TemplateGenerator] Modo: {fonte}")

    return gen.gerar(tipo, dados)