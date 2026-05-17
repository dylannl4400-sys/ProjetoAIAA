"""
geracao/views.py

Dois modos de geração:
    POST /geracao/gerar/         → modo simples (sem RAG)
    POST /geracao/gerar-rag/     → modo RAG (com acórdãos indexados)
"""
import json, os, re, tempfile
from django.http  import JsonResponse, FileResponse
from django.views.decorators.csrf  import csrf_exempt
from django.views.decorators.http  import require_POST
from aiaa_init import get_llm, get_store

# Templates com formulários fixos no frontend
TEMPLATES_BUILTIN = {"contestacao", "peticao", "requerimento", "cessacao"}


def _extrair_marcadores_docx(caminho: str) -> list:
    """Lê um ficheiro .docx e devolve todos os marcadores {{...}} encontrados."""
    try:
        from docx import Document
        doc = Document(caminho)
        marcadores = set()

        def _scan_paragrafos(paragrafos):
            for para in paragrafos:
                found = re.findall(r'\{\{([^}]+)\}\}', para.text)
                marcadores.update(found)

        _scan_paragrafos(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _scan_paragrafos(cell.paragraphs)

        return sorted(marcadores)
    except Exception:
        return []


def _validar(body: dict) -> str | None:
    """Devolve mensagem de erro ou None se OK."""
    from geracao.models import Template
    if not Template.objects.filter(identificador=body.get("template"), ativo=True).exists():
        return f"Template '{body.get('template')}' não encontrado ou inactivo."
    # Para templates builtin, exigir campos mínimos
    if body.get("template") in TEMPLATES_BUILTIN:
        if not body.get("nome_parte_a", "").strip():
            return "nome_parte_a obrigatório."
        if not body.get("descricao_factos", "").strip():
            return "descricao_factos obrigatório."
    return None


def _gerar_e_devolver(body: dict, usar_rag: bool):
    """Lógica comum a ambas as views."""
    erro = _validar(body)
    if erro:
        return JsonResponse({"erro": erro}, status=400)

    tipo = body["template"]

    try:
        from template_generator import TemplateGenerator

        from geracao.models import Template
        template_obj = Template.objects.get(identificador=tipo)

        # Se usar_rag é False, não passamos LLM para fazer apenas substituição direta
        gen = TemplateGenerator(
            llm   = get_llm() if usar_rag else None,
            store = get_store() if usar_rag else None,
        )

        conteudo = gen._gerar_conteudo(tipo, body)
        caminho  = gen.gerar(template_obj.nome_ficheiro_docx, body, conteudo=conteudo)

        # Guardar registo na BD
        try:
            from geracao.models import PecaGerada
            PecaGerada.objects.create(
                template    = template_obj,
                trabalhador = body.get("nome_parte_a", ""),
                processo    = body.get("referencia_processo", ""),
            )
        except Exception:
            pass

        # Headers extra para info sobre RAG
        acordaos_usados = conteudo.get("acordaos", [])
        nome_ficheiro   = f"{tipo}_{body.get('nome_parte_a','documento').replace(' ','_')}.docx"

        resp = FileResponse(
            open(caminho, "rb"),
            as_attachment=True,
            filename=nome_ficheiro,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        resp["X-Acordaos-RAG"] = str(len(acordaos_usados))
        resp["X-Fonte"]        = conteudo.get("fonte", "?")
        return resp

    except FileNotFoundError as e:
        return JsonResponse({"erro": f"Template não encontrado: {e}"}, status=404)
    except ImportError as e:
        return JsonResponse({"erro": f"pip install python-docx ({e})"}  , status=500)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({"erro": str(e)}, status=500)


@csrf_exempt
@require_POST
def gerar_peca(request):
    """POST /geracao/gerar/"""
    try:
        body = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({"erro": "JSON inválido."}, status=400)
    return _gerar_e_devolver(body, usar_rag=False)


@csrf_exempt
@require_POST
def gerar_peca_rag(request):
    """POST /geracao/gerar-rag/"""
    try:
        body = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({"erro": "JSON inválido."}, status=400)
    return _gerar_e_devolver(body, usar_rag=True)


def templates_listar(request):
    """GET /geracao/templates/"""
    from geracao.models import Template
    templates = Template.objects.filter(ativo=True)
    return JsonResponse({"templates": [
        {
            "id_db":  t.id,
            "id":     t.identificador,
            "nome":   t.nome_exibicao,
            "descricao": t.descricao,
            "campos": t.campos,   # lista de marcadores detectados no docx
        }
        for t in templates
    ]})


@csrf_exempt
@require_POST
def template_upload(request):
    """
    POST /geracao/templates/upload/
    Multipart form:
        nome: str
        identificador: str
        descricao: str (opcional)
        ficheiro: file (.docx)
    """
    nome          = request.POST.get("nome")
    identificador = request.POST.get("identificador")
    descricao     = request.POST.get("descricao", "")
    ficheiro      = request.FILES.get("ficheiro")

    if not nome or not identificador or not ficheiro:
        return JsonResponse({"erro": "Campos obrigatórios em falta (nome, identificador ou ficheiro)."}, status=400)

    if not ficheiro.name.lower().endswith(".docx"):
        return JsonResponse({"erro": "Apenas ficheiros .docx são permitidos."}, status=400)

    # Pasta de destino
    destino_pasta = os.path.join("pipeline", "legal_docs")
    os.makedirs(destino_pasta, exist_ok=True)

    nome_ficheiro = f"template_{identificador}.docx"
    caminho = os.path.join(destino_pasta, nome_ficheiro)

    try:
        with open(caminho, "wb+") as destination:
            for chunk in ficheiro.chunks():
                destination.write(chunk)

        # Extrair marcadores do documento
        campos_detectados = _extrair_marcadores_docx(caminho)

        from geracao.models import Template
        Template.objects.update_or_create(
            identificador=identificador,
            defaults={
                "nome_exibicao":    nome,
                "nome_ficheiro_docx": nome_ficheiro,
                "descricao":        descricao,
                "campos":           campos_detectados,
                "ativo":            True,
            }
        )
        return JsonResponse({
            "sucesso":     True,
            "identificador": identificador,
            "campos":      campos_detectados,
        })
    except Exception as e:
        return JsonResponse({"erro": str(e)}, status=500)


@csrf_exempt
@require_POST
def template_apagar(request, tid):
    """POST /geracao/templates/<tid>/apagar/"""
    from geracao.models import Template
    try:
        t = Template.objects.get(id=tid)
        t.ativo = False  # Soft delete
        t.save()
        return JsonResponse({"sucesso": True})
    except Template.DoesNotExist:
        return JsonResponse({"erro": "Template não encontrado."}, status=404)